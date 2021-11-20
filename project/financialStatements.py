# -*- coding: UTF-8 -*-

import pandas as pd
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm

from project.settings import setStyle
from project.utils import createBorderedTable,setCell,addLandscapeContent,searchModel,getNoteNum,to_chinese,\
    set_cell_border,checkLeftSpace,combineTitles


small = "tableSmallCharacter"
smaller = "tableSmallerCharacter"

# 获取列标题前部空格数量和排版格式
def getAlignAndText(alignType,originText):
    if alignType=="center":
        return originText,WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        return  " "*4*alignType+originText,WD_PARAGRAPH_ALIGNMENT.LEFT

# 过滤报表空行数据
# 过滤资产负债表、利润表、现金流量表数据
def FilterFsNewRecords(display,newRecords):
    resRecords = []
    for key in range(0, len(newRecords)):
        if display:
            if newRecords[key]["display"]:
                resRecords.append(newRecords[key])
            else:
                if abs(newRecords[key]["startDate"]) > 1e-6 or abs(newRecords[key]["endDate"]) > 1e-6:
                    resRecords.append(newRecords[key])
        else:
            return newRecords
    return resRecords
# 过滤所有者权益变动标数据
def FilterOsNewRecords(display,newRecords):
    '''

    :param display: 是否必须显示
    :param newRecords: 数据集合
    :return:
    '''
    resRecords = []
    for key in range(0, len(newRecords)):
        if display:
            if newRecords[key]["display"]:
                resRecords.append(newRecords[key])
            else:
                if abs(newRecords[key]["paidInCapital"]) > 1e-6 or \
                        abs(newRecords[key]["preferedStock"]) > 1e-6 or \
                        abs(newRecords[key]["perpetualDebt"]) > 1e-6 or \
                        abs(newRecords[key]["otherEquityInstruments"])> 1e-6 or \
                        abs(newRecords[key]["capitalReserve"]) > 1e-6 or \
                        abs(newRecords[key]["treasuryStock"]) > 1e-6 or \
                        abs(newRecords[key]["otherComprehensiveIncome"]) > 1e-6 or \
                        abs(newRecords[key]["specialReserve"]) > 1e-6 or \
                        abs(newRecords[key]["surplusReserve"]) > 1e-6 or \
                        abs(newRecords[key]["generalRiskReserve"]) > 1e-6 or \
                        abs(newRecords[key]["undistributedProfit"]) > 1e-6 or \
                        abs(newRecords[key]["subtotal"]) > 1e-6 or \
                        abs(newRecords[key]["minorityInterests"]) > 1e-6 or \
                        abs(newRecords[key]["totalOwnerEquity"]) > 1e-6:
                    resRecords.append(newRecords[key])
        else:
            return newRecords
    return resRecords


# 添加表头：如资产负债表
def addTableTitle(document,name):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run(name,style="first")
# 添加表头下一行，表格前一行，如：单位：***公司 单位：元
def addFirstLine(document,context):

    # 公司名称
    companyName = context["report_params"]["companyName"]
    # 报告日期
    reportDate = context["report_params"]["reportDate"]
    # 人民币单位
    currencyUnit = context["notes_params"]["currencyUnit"]

    table = document.add_table(1, 3)
    table.rows[0].height_rule  =WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Cm(0.8)
    table.cell(0, 0).width = Cm(20)
    table.cell(0, 1).width = Cm(10)
    table.cell(0, 2).width = Cm(10)
    setCell(table.cell(0, 0), "编制单位：{}".format(companyName), WD_PARAGRAPH_ALIGNMENT.LEFT, toFloat=True, style="tableSmallCharacter")
    setCell(table.cell(0, 1), reportDate, WD_PARAGRAPH_ALIGNMENT.CENTER, toFloat=True, style="tableSmallCharacter")
    setCell(table.cell(0, 2), "单位：{}".format(currencyUnit), WD_PARAGRAPH_ALIGNMENT.RIGHT, toFloat=True, style="tableSmallCharacter")
# 添加最后一行 向报表中添加最后一行。如：会计机构负责人：  财务负责人：
def add_last_line(document):
    table = document.add_table(1, 3)
    table.cell(0, 0).width = Cm(10)
    table.cell(0, 1).width = Cm(13)
    table.cell(0, 2).width = Cm(10)
    setCell(table.cell(0, 0), "法定代表人：", WD_PARAGRAPH_ALIGNMENT.LEFT, toFloat=True,
            style="tableSmallCharacter")
    setCell(table.cell(0, 1), "主管会计工作的负责人：", WD_PARAGRAPH_ALIGNMENT.LEFT, toFloat=True, style="tableSmallCharacter")
    setCell(table.cell(0, 2), "会计机构负责人：", WD_PARAGRAPH_ALIGNMENT.LEFT, toFloat=True,
            style="tableSmallCharacter")
# 添加标题
# 设置国有企业报表标题
def setFsHeaderState(table,titles):
    table.cell(0, 0).width = Cm(13)
    table.cell(0, 1).width = Cm(2)
    table.cell(0, 2).width = Cm(5)
    table.cell(0, 3).width = Cm(5)
    table.cell(0, 4).width = Cm(5)
    setCell(table.cell(0, 0), titles[0], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 1), titles[1], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 2), titles[2], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 3), titles[3], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 4), titles[4], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
# 设置国有企业所有者权益表标题


# 设置国有企业所有者权益变动表表头
def setOwnerHeaderState(table,period):
    table.cell(0, 0).width = Cm(8)
    table.cell(1, 0).width = Cm(8)
    table.cell(2, 0).width = Cm(8)
    table.cell(3, 0).width = Cm(8)
    table.cell(0, 1).width = Cm(2)
    table.cell(1, 1).width = Cm(2)
    table.cell(2, 1).width = Cm(2)
    table.cell(3, 1).width = Cm(2)

    values = [
        ["项            目","行次",period,"nan","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan"],
        ["nan","nan","归属于母公司所有者权益","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan","少数股东权益","所有者权益合计"],
        ["nan","nan","实收资本(或股本)","其他权益工具","nan","nan","资本公积","减:库存股","其他综合收益","专项储备","盈余公积","△一般风险准备","未分配利润","小计","nan","nan"],
        ["nan","nan","nan","优先股","永续债","其他","nan","nan","nan","nan","nan","nan","nan","nan","nan","nan"],
    ]
    for i,row in enumerate(values):
        for j,value in enumerate(row):
            setCell(table.cell(i, j), value, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})

# 设置上市公司报表标题
def setFsHeaderList(table,titles):
    table.cell(0, 0).width = Cm(13)
    table.cell(0, 1).width = Cm(5)
    table.cell(0, 2).width = Cm(5)
    table.cell(0, 3).width = Cm(5)
    setCell(table.cell(0, 0), titles[0], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 1), titles[1], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 2), titles[2], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(0, 3), titles[3], WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
# 设置上市公司合并所有者权益表标题
def setOwnerHeaderList(table,period):
    table.cell(0, 0).width = Cm(8)
    table.cell(1, 0).width = Cm(8)
    table.cell(2, 0).width = Cm(8)
    table.cell(3, 0).width = Cm(8)
    table.cell(0, 1).width = Cm(2)
    table.cell(1, 1).width = Cm(2)
    table.cell(2, 1).width = Cm(2)
    table.cell(3, 1).width = Cm(2)

    values = [
        ["项            目",  period, "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan",
         "nan", "nan"],
        ["nan", "归属于母公司所有者权益", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan",
         "少数股东权益", "所有者权益合计"],
        ["nan","实收资本(或股本)", "其他权益工具", "nan", "nan", "资本公积", "减:库存股", "其他综合收益", "专项储备", "盈余公积", "一般风险准备",
         "未分配利润", "小计", "nan", "nan"],
        ["nan",  "nan", "优先股", "永续债", "其他", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan"],
    ]

    for i,row in enumerate(values):
        for j,value in enumerate(row):
            setCell(table.cell(i, j), value, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})
 # 设置上市公司单体所有者权益表标题
def setOwnerHeaderListSingle(table, period):
    table.cell(0, 0).width = Cm(8)
    table.cell(1, 0).width = Cm(8)
    table.cell(2, 0).width = Cm(8)
    table.cell(0, 1).width = Cm(2)
    table.cell(1, 1).width = Cm(2)
    table.cell(2, 1).width = Cm(2)

    values = [
        ["项            目", period, "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan", "nan"],
        ["nan", "实收资本(或股本)", "其他权益工具", "nan", "nan", "资本公积", "减:库存股", "其他综合收益", "专项储备", "盈余公积",
         "未分配利润", "所有者权益合计"],
        ["nan", "nan", "优先股", "永续债", "其他", "nan", "nan", "nan", "nan", "nan", "nan", "nan"],
    ]
    for i,row in enumerate(values):
        for j,value in enumerate(row):
            setCell(table.cell(i, j), value, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(0, 0), "项            目", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(0, 1), period, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 1), "实收资本(或股本)", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 2), "其他权益工具", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 5), "资本公积", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 6), "减:库存股", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 7), "其他综合收益", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 8), "专项储备", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 9), "盈余公积", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 10), "未分配利润", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(1, 11), "所有者权益合计", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(2, 2), "优先股", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(2, 3), "永续债", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # setCell(table.cell(2, 4), "其他", WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    # # 合并单元格
    # # 第一列
    # for i in range(1, 3):
    #     table.cell(i, 0).merge(table.cell(i - 1, 0))
    # # 第一行
    # for i in range(2, 12):
    #     table.cell(0, i).merge(table.cell(0, i - 1))
    # # 第二列
    # table.cell(2, 1).merge(table.cell(1, 1))
    # # 第二行
    # for i in range(3, 5):
    #     table.cell(1, i).merge(table.cell(1, i - 1))
    # # 第七列到十二列
    # for i in range(5, 12):
    #     table.cell(2, i).merge(table.cell(1, i))

    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})

# 添加一行记录
# 向国有企业报表中添加一行记录
def addFsTableLineState(table,key,newRecords,num):
    smallNum = to_chinese(newRecords[key]["noteNum"])
    if smallNum==0:
        recordNum = ""
    else:
        recordNum = "{}({})".format(num,smallNum)

    table.cell(key, 0).width = Cm(13)
    table.cell(key, 1).width = Cm(2)
    table.cell(key, 2).width = Cm(5)
    table.cell(key, 3).width = Cm(5)
    table.cell(key, 4).width = Cm(5)
    setCell(table.cell(key, 0), *getAlignAndText(newRecords[key]["type"], newRecords[key]["name"]), False,
            smaller)
    setCell(table.cell(key, 1), key, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(key, 2), newRecords[key]["endDate"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key, 3), newRecords[key]["startDate"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key, 4),recordNum, WD_PARAGRAPH_ALIGNMENT.CENTER, True, smaller)
# 向上市公司报表中添加一行记录
def addFsTableLineList(table,key,newRecords,num):
    smallNum = to_chinese(newRecords[key]["noteNum"])
    if smallNum == 0:
        recordNum = ""
    else:
        recordNum = "{}({})".format(num, smallNum)

    table.cell(key, 0).width = Cm(13)
    table.cell(key, 1).width = Cm(5)
    table.cell(key, 2).width = Cm(5)
    table.cell(key, 3).width = Cm(5)
    setCell(table.cell(key, 0), *getAlignAndText(newRecords[key]["type"], newRecords[key]["name"]), False,
            smaller)
    setCell(table.cell(key, 1), recordNum, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(key, 2), newRecords[key]["endDate"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key, 3), newRecords[key]["startDate"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
# 向国有企业所有者权益变动表中添加一行记录
def addOsTableLineState(table,key,newRecords):
    table.cell(key + 3, 0).width = Cm(8)
    table.cell(key + 3, 1).width = Cm(1)

    setCell(table.cell(key + 3, 0), *getAlignAndText(newRecords[key]["type"], newRecords[key]["name"]), False,
            smaller)
    setCell(table.cell(key + 3, 1), key, WD_PARAGRAPH_ALIGNMENT.CENTER, False, smaller)
    setCell(table.cell(key + 3, 2), newRecords[key]["paidInCapital"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 3), newRecords[key]["preferedStock"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 4), newRecords[key]["perpetualDebt"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 5), newRecords[key]["otherEquityInstruments"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 6), newRecords[key]["capitalReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 7), newRecords[key]["treasuryStock"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 8), newRecords[key]["otherComprehensiveIncome"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 9), newRecords[key]["specialReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 10), newRecords[key]["surplusReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 11), newRecords[key]["generalRiskReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 12), newRecords[key]["undistributedProfit"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 13), newRecords[key]["subtotal"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 14), newRecords[key]["minorityInterests"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + 3, 15), newRecords[key]["totalOwnerEquity"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
# 向上市公司所有者权益变动表中添加一行记录
def addOsTableLineList(table,key,newRecords,reportType):

    if reportType=="合并":
        step = 3
    else:
        step = 2
    table.cell(key + step, 0).width = Cm(8)
    table.cell(key + step, 1).width = Cm(1)

    setCell(table.cell(key + step, 0), *getAlignAndText(newRecords[key]["type"], newRecords[key]["name"]), False,
            smaller)
    setCell(table.cell(key + step, 1), newRecords[key]["paidInCapital"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 2), newRecords[key]["preferedStock"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 3), newRecords[key]["perpetualDebt"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 4), newRecords[key]["otherEquityInstruments"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 5), newRecords[key]["capitalReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 6), newRecords[key]["treasuryStock"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 7), newRecords[key]["otherComprehensiveIncome"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 8), newRecords[key]["specialReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    setCell(table.cell(key + step, 9), newRecords[key]["surplusReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
            smaller)
    if reportType == "合并":
        setCell(table.cell(key + step, 10), newRecords[key]["generalRiskReserve"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
        setCell(table.cell(key + step, 11), newRecords[key]["undistributedProfit"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
        setCell(table.cell(key + step, 12), newRecords[key]["subtotal"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
        setCell(table.cell(key + step, 13), newRecords[key]["minorityInterests"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
        setCell(table.cell(key + step, 14), newRecords[key]["totalOwnerEquity"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
    else:
        setCell(table.cell(key + step, 10), newRecords[key]["undistributedProfit"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)
        setCell(table.cell(key + step, 11), newRecords[key]["subtotal"], WD_PARAGRAPH_ALIGNMENT.RIGHT, True,
                smaller)


# 获取记录中全部为0的列编号
def get_zero_column_num(records):
    zero_column_num = []
    contrast = {

    }
    df = pd.DataFrame(records)
    newdf = df.fillna(value=0.00)
    for key,i in enumerate(newdf):
        if np.issubdtype(newdf[i].dtype, np.number):
            if (newdf[i].abs()<1e-7).all():
                zero_column_num.append(key-1)
    return zero_column_num

# 添加国有企业报表
def addFinancialStatementsState(document,name,titles,oldRecords,context,lastSection=True,display=True):
    # 报告类型
    reportType = context["report_params"]["type"]
    isParent = False
    if reportType == "合并":
        if "合并" not in name:
            isParent = True
    num = getNoteNum(context, isParent)

    addTableTitle(document,name)
    addFirstLine(document,context)
    document.add_section(start_type=0)
    newRecords = FilterFsNewRecords(display,oldRecords)
    table = createBorderedTable(document,len(newRecords),5,innerLine="single")
    table.columns[0].width = Cm(7)
    setFsHeaderState(table,titles)

    for key in range(1,len(newRecords)):
        addFsTableLineState(table, key, newRecords,num)

    document.add_section(start_type=0)
    add_last_line(document)
    if lastSection:
        document.add_section()
# 添加上市公司报表
def addFinancialStatementsList(document,name,titles,oldRecords,context,lastSection=True,display=True):
    # 报告类型
    reportType = context["report_params"]["type"]
    isParent = False
    if reportType=="合并":
        if "合并" not in name:
            isParent = True
    num = getNoteNum(context,isParent)

    addTableTitle(document,name)
    addFirstLine(document,context)
    document.add_section(start_type=0)
    newRecords = FilterFsNewRecords(display,oldRecords)
    table = createBorderedTable(document,len(newRecords),4,innerLine="single")
    table.columns[0].width = Cm(7)
    setFsHeaderList(table,titles)

    for key in range(1,len(newRecords)):
        addFsTableLineList(table, key, newRecords,num)

    document.add_section(start_type=0)
    add_last_line(document)
    if lastSection:
        document.add_section()
# 删除docx中表格中的列
def delete_columns(table, columns):
    # sort columns descending
    columns.sort(reverse=True)

    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for ci in columns:
        for cell in table.column_cells(ci):
            cell._tc.getparent().remove(cell._tc)

        # Delete column reference.
        col_elem = grid[ci]
        grid.remove(col_elem)

# 获取所有者权益变动表标题头
def get_os_table_titles(table,titleRowNum):
    titles = []
    for i in range(titleRowNum):
        row = table.rows[i]
        rowTitle = []
        for cell in row.cells:
            rowTitle.append(cell.text)
            # setCell(cell, "nan", WD_PARAGRAPH_ALIGNMENT.CENTER)
        titles.append(rowTitle)
    return titles
# 添加国有企业所有者权益变动表
def addOnwerEquityState(document,name,oldRecords,context,period,display=True):
    addTableTitle(document,name)
    addFirstLine(document,context)
    document.add_section(start_type=0)
    newRecords = FilterOsNewRecords(display,oldRecords)
    table = createBorderedTable(document,len(newRecords)+3,16,innerLine="single")
    # 设置表头内容
    setOwnerHeaderState(table, period)

    for key in range(1,len(newRecords)):
        addOsTableLineState(table, key, newRecords)
    # 删除全部为0的列
    zero_column_num = get_zero_column_num(newRecords)
    delete_columns(table,zero_column_num)
    # 删除空白行标题，标题第四行
    row3 = table.rows[3]
    cell3Text = [cell.text for cell in row3.cells]
    # 列标题长度
    column_length = len(table.columns)
    if len(set(cell3Text))==1 and list(set(cell3Text))[0]=="":
        row3._element.getparent().remove(row3._element)
        titles = get_os_table_titles(table,3)
        combineTitles(titles,table,[[2,column_length-2],[2,column_length-1]],lastLine=False)
    else:
        titles = get_os_table_titles(table, 4)
        combineTitles(titles, table, [[2, column_length - 2], [2, column_length - 1]],lastLine=True)

    document.add_section(start_type=0)
    add_last_line(document)


# 添加上市公司所有者权益变动表
def addOnwerEquityList(document,name,oldRecords,context,period,reportType,display=True):
    addTableTitle(document,name)
    addFirstLine(document,context)
    document.add_section(start_type=0)
    newRecords = FilterOsNewRecords(display,oldRecords)
    if reportType == "合并":
        table = createBorderedTable(document,len(newRecords)+3,15,innerLine="single")
        setOwnerHeaderList(table,period)
        for key in range(1, len(newRecords)):
            addOsTableLineList(table, key, newRecords, reportType)
        # 删除全部为0的列
        zero_column_num = get_zero_column_num(newRecords)
        delete_columns(table, zero_column_num)
        # 删除空白行标题，标题第四行
        row3 = table.rows[3]
        cell3Text = [cell.text for cell in row3.cells]
        # 列标题长度
        column_length = len(table.columns)
        if len(set(cell3Text)) == 1 and list(set(cell3Text))[0] == "":
            row3._element.getparent().remove(row3._element)
            titles = get_os_table_titles(table, 3)
            combineTitles(titles, table, [[2, column_length - 2], [2, column_length - 1]],lastLine=False)
        else:
            titles = get_os_table_titles(table, 4)
            combineTitles(titles, table, [[2, column_length - 2], [2, column_length - 1]],lastLine=True)

    else:
        table = createBorderedTable(document, len(newRecords) + 2, 12, innerLine="single")
        setOwnerHeaderListSingle(table, period)
        for key in range(1, len(newRecords)):
            addOsTableLineList(table, key, newRecords, reportType)
        # 删除全部为0的列
        zero_column_num = get_zero_column_num(newRecords)
        delete_columns(table, zero_column_num)
        # 删除空白行标题，标题第三行
        row2 = table.rows[2]
        cell2Text = [cell.text for cell in row2.cells]
        # 列标题长度
        column_length = len(table.columns)
        if len(set(cell2Text)) == 1 and list(set(cell2Text))[0] == "":
            row2._element.getparent().remove(row2._element)
            titles = get_os_table_titles(table, 2)
            combineTitles(titles, table,[],lastLine=False)
        else:
            titles = get_os_table_titles(table, 3)
            combineTitles(titles, table, [],lastLine=True)

    document.add_section(start_type=0)
    add_last_line(document)

# 添加所有者权益变动表
def addOwnership(document,context,companyType,reportType,ownerRecordsCombineThis,ownerRecordsCombineLast,ownerRecordsSingleThis,
                 ownerRecordsSingleLast):
    if companyType == "上市公司":
        if reportType == "合并":
            addOnwerEquityList(document,"合并所有者权益变动表", ownerRecordsCombineThis,context,"本 期 金 额","合并")
            document.add_section()
            addOnwerEquityList(document,"合并所有者权益变动表（续）", ownerRecordsCombineLast, context,"上 期 金 额","合并")
            document.add_section()
            addOnwerEquityList(document, "所有者权益变动表", ownerRecordsSingleThis, context, "本 期 金 额","单体")
            document.add_section()
            addOnwerEquityList(document, "所有者权益变动表（续）", ownerRecordsSingleLast, context, "上 期 金 额","单体")
        else:
            addOnwerEquityList(document, "所有者权益变动表", ownerRecordsSingleThis, context, "本 期 金 额","单体")
            document.add_section()
            addOnwerEquityList(document, "所有者权益变动表（续）", ownerRecordsSingleLast, context, "上 期 金 额","单体")
    elif companyType == "国有企业":
        if reportType == "合并":
            addOnwerEquityState(document,"合并所有者权益变动表", ownerRecordsCombineThis, context,"本 期 金 额")
            document.add_section()
            addOnwerEquityState(document,"合并所有者权益变动表（续）", ownerRecordsCombineLast, context,"上 期 金 额")
            document.add_section()
            addOnwerEquityState(document, "所有者权益变动表", ownerRecordsSingleThis, context, "本 期 金 额")
            document.add_section()
            addOnwerEquityState(document, "所有者权益变动表（续）", ownerRecordsSingleLast, context, "上 期 金 额")
        else:
            addOnwerEquityState(document, "所有者权益变动表", ownerRecordsSingleThis, context, "本 期 金 额")
            document.add_section()
            addOnwerEquityState(document, "所有者权益变动表（续）", ownerRecordsSingleLast, context, "上 期 金 额")
# 添加报表
def reportForm(document,context,balanceTitles,profitTitles,
               assetsRecordsCombine,liabilitiesRecordsCombine,profitRecordsCombine,cashRecordsCombine,assetsRecordsSingle,
               liabilitiesRecordsSingle,profitRecordsSingle,cashRecordsSingle,ownerRecordsCombineThis,ownerRecordsCombineLast,
               ownerRecordsSingleThis,ownerRecordsSingleLast):
    # 公司类型
    companyType = context["report_params"]["companyType"]
    # 公司名称
    companyName = context["report_params"]["companyName"]
    # 报告类型
    reportType = context["report_params"]["type"]
    # 报告日期
    reportDate = context["report_params"]["reportDate"]
    # 获取报告起始日
    startYear = reportDate[:4]
    # 报告期间
    reportPeriod = context["report_params"]["reportPeriod"]
    lastPeriod = str.replace(reportPeriod, reportPeriod[:4], str(int(reportPeriod[:4]) - 1))
    # 人民币单位
    currencyUnit = context["notes_params"]["currencyUnit"]

    if companyType=="国有企业":
        if reportType == "合并":
            # 添加合并资产负债表
            addFinancialStatementsState(document,"合并资产负债表", balanceTitles, assetsRecordsCombine, context)
            addFinancialStatementsState(document,"合并资产负债表(续)", balanceTitles, liabilitiesRecordsCombine, context)
            addFinancialStatementsState(document,"资产负债表", balanceTitles, assetsRecordsSingle, context)
            addFinancialStatementsState(document,"资产负债表(续)", balanceTitles, liabilitiesRecordsSingle, context)
            addFinancialStatementsState(document,"合并利润表", profitTitles, profitRecordsCombine, context)
            addFinancialStatementsState(document,"利润表", profitTitles, profitRecordsSingle, context)
            addFinancialStatementsState(document,"合并现金流量表", profitTitles, cashRecordsCombine, context)
            addFinancialStatementsState(document,"现金流量表", profitTitles, cashRecordsSingle, context, lastSection=False)

        else:
            addFinancialStatementsState(document,"资产负债表", balanceTitles, assetsRecordsSingle, context)
            addFinancialStatementsState(document,"资产负债表(续)", balanceTitles, liabilitiesRecordsSingle,context)
            addFinancialStatementsState(document,"利润表", profitTitles, profitRecordsSingle, context)
            addFinancialStatementsState(document,"现金流量表", profitTitles, cashRecordsSingle, context,lastSection=False)

    elif companyType == "上市公司":
        if reportType == "合并":
            addFinancialStatementsList(document,"合并资产负债表", balanceTitles, assetsRecordsCombine, context)
            addFinancialStatementsList(document,"合并资产负债表(续)", balanceTitles, liabilitiesRecordsCombine, context)
            addFinancialStatementsList(document,"资产负债表", balanceTitles, assetsRecordsSingle, context)
            addFinancialStatementsList(document,"资产负债表(续)", balanceTitles, liabilitiesRecordsSingle, context)
            addFinancialStatementsList(document,"合并利润表", profitTitles, profitRecordsCombine, context)
            addFinancialStatementsList(document,"利润表", profitTitles, profitRecordsSingle, context)
            addFinancialStatementsList(document,"合并现金流量表", profitTitles, cashRecordsCombine, context)
            addFinancialStatementsList(document,"现金流量表", profitTitles, cashRecordsSingle, context,lastSection=False)

        else:
            addFinancialStatementsList(document,"资产负债表", balanceTitles, assetsRecordsSingle, context)
            addFinancialStatementsList(document,"资产负债表(续)", balanceTitles, liabilitiesRecordsSingle, context)
            addFinancialStatementsList(document,"利润表", profitTitles, profitRecordsSingle, context)
            addFinancialStatementsList(document,"现金流量表", profitTitles, cashRecordsSingle, context,lastSection=False)


    else:
        pass
    # 添加所有者权益变动表
    addLandscapeContent(document, addOwnership,  context,companyType, reportType,
                        ownerRecordsCombineThis,ownerRecordsCombineLast,ownerRecordsSingleThis,ownerRecordsSingleLast)



def addFs(document,context,comparativeTable,balanceTitlesState,balanceTitlesList,profitTitlesState,profitTitlesList):
    # 公司类型：上市公司、国有企业
    companyType = context["report_params"]["companyType"]
    # 获取报表数据
    assetsRecordsCombine = searchModel(companyType,"合并","资产表",comparativeTable)
    liabilitiesRecordsCombine = searchModel(companyType,"合并","负债表",comparativeTable)
    profitRecordsCombine = searchModel(companyType,"合并","利润表",comparativeTable)
    cashRecordsCombine = searchModel(companyType,"合并","现金流量表",comparativeTable)
    assetsRecordsSingle = searchModel(companyType,"单体","资产表",comparativeTable)
    liabilitiesRecordsSingle = searchModel(companyType,"单体","负债表",comparativeTable)
    profitRecordsSingle = searchModel(companyType,"单体","利润表",comparativeTable)
    cashRecordsSingle = searchModel(companyType,"单体","现金流量表",comparativeTable)
    ownerRecordsCombineThis = searchModel(companyType,"合并","本期所有者权益变动表",comparativeTable)
    ownerRecordsCombineLast = searchModel(companyType,"合并","上期所有者权益变动表",comparativeTable)
    ownerRecordsSingleThis = searchModel(companyType,"单体","本期所有者权益变动表",comparativeTable)
    ownerRecordsSingleLast = searchModel(companyType,"单体","上期所有者权益变动表",comparativeTable)

    if companyType=="上市公司":
        reportForm(document,context,balanceTitlesList,profitTitlesList,assetsRecordsCombine,
                   liabilitiesRecordsCombine, profitRecordsCombine, cashRecordsCombine, assetsRecordsSingle,
                   liabilitiesRecordsSingle, profitRecordsSingle, cashRecordsSingle, ownerRecordsCombineThis, ownerRecordsCombineLast,
                   ownerRecordsSingleThis, ownerRecordsSingleLast)
    else:
        reportForm(document,context,balanceTitlesState,profitTitlesState,assetsRecordsCombine,
                   liabilitiesRecordsCombine, profitRecordsCombine, cashRecordsCombine, assetsRecordsSingle,
                   liabilitiesRecordsSingle, profitRecordsSingle, cashRecordsSingle, ownerRecordsCombineThis, ownerRecordsCombineLast,
                   ownerRecordsSingleThis, ownerRecordsSingleLast)


def test():
    from project.data import testcontext
    from project.constants import comparativeTable, tables, contrastSubjects, CURRENTPATH, PARENTPATH, \
        balanceTitlesState, balanceTitlesList, profitTitlesState, profitTitlesList
    from project.fsmodel import fillTable
    from project.computeNo import computeNo

    CURRENTPATH = r"E:\审计\我的文件2021\义务2020年审\义乌市粮食收储有限公司TB及附注\0义乌市粮食收储有限公司.xlsx"
    PARENTPATH = r"E:\审计\我的文件2021\义务2020年审\义乌市粮食收储有限公司TB及附注\0义乌市粮食收储有限公司.xlsx"

    document = Document()
    # 设置中文标题
    setStyle(document)
    # 填充报表数据
    fillTable(testcontext, comparativeTable, tables, contrastSubjects, CURRENTPATH, PARENTPATH)
    # 计算附注编码
    computeNo(testcontext, comparativeTable)
    addFs(document, testcontext, comparativeTable,  balanceTitlesState, balanceTitlesList,profitTitlesState, profitTitlesList)

    document.save("fs.docx")

if __name__ == '__main__':
    test()
