# -*- coding: UTF-8 -*-

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from project.settings import setStyle
from project.utils import add_page_number


# 获取公司称谓
def getAppellation(companyName):
    if "股份" in companyName:
        appellation = companyName + "全体股东："
    else:
        appellation = companyName + "董事会："
    return appellation


# 根据会计师事务所名称长度调整会计师事务所显示
def getAccountFirmNames(accountFirm):
    res = []
    if len(accountFirm) < 16:
        res.append(accountFirm)
    else:
        cnBrace = accountFirm.rfind("（")
        enBrace = accountFirm.rfind("(")
        if cnBrace != -1:
            res.append(accountFirm[:cnBrace])
            res.append(accountFirm[cnBrace:])
        else:
            if enBrace != -1:
                res.append(accountFirm[:enBrace])
                res.append(accountFirm[enBrace:])
            else:
                per = int(len(accountFirm) / 2)
                res.append(accountFirm[:per])
                res.append(accountFirm[per:])
    return res


# 添加其他信息
def addOtherInfo(document,context):
    paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph4.add_run(
        "{}管理层（以下简称管理层）对其他信息负责。其他信息包括年度报告中涵盖的信息，但不包括财务报表和我们的审计报告。".format(context["report_params"]["CompanyAbbrName"]),
        style="zero")
    paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph4.add_run("我们对财务报表发表的审计意见不涵盖其他信息，我们也不对其他信息发表任何形式的鉴证结论。", style="zero")
    paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph4.add_run("结合我们对财务报表的审计，我们的责任是阅读其他信息，在此过程中，考虑其他信息是否与财务报表或我们在审计过程中了解到的情况存在重大不一致或者似乎存在重大错报。", style="zero")
    paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph4.add_run("基于我们已执行的工作，如果我们确定其他信息存在重大错报，我们应当报告该事实。在这方面，我们无任何事项需要报告。",
                       style="zero")


# 添加管理层责任
def addManagementResponsebility(document,context):
    companyAbbrName = context["report_params"]["CompanyAbbrName"]
    responsibility1 = [companyAbbrName, "管理层负责按照企业会计准则的规定编制财务报表，使其实现公允反映，并设计、执行和维护必要的内部控制，以使财务报表不存在由于舞弊或错误导致的重大错报。"]
    paragraph7 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph7.add_run("".join(responsibility1), style="zero")

    responsibility2 = [
        "在编制财务报表时，管理层负责评估" + companyAbbrName + "的持续经营能力，披露与持续经营相关的事项（如适用），并运用持续经营假设，除非计划进行清算、终止运营或别无其他现实的选择。"]
    paragraph8 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph8.add_run("".join(responsibility2), style="zero")

    responsibility3 = ["治理层负责监督", companyAbbrName, "的财务报告过程。"]
    paragraph9 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph9.add_run("".join(responsibility3), style="zero")


# 添加注册会计师责任
def addCpaResponsibility(document,context):
    companyAbbrName = context["report_params"]["CompanyAbbrName"]
    paragraph11 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph11.add_run(
        "我们的目标是对财务报表整体是否不存在由于舞弊或错误导致的重大错报获取合理保证，并出具包含审计意见的审计报告。合理保证是高水平的保证，但并不能保证按照审计准则执行的审计在某一重大错报存在时总能发现。错报可能由于舞弊或错误导致，如果合理预期错报单独或汇总起来可能影响财务报表使用者依据财务报表作出的经济决策，则通常认为错报是重大的。",
        style="zero")

    paragraph12 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph12.add_run("在按照审计准则执行审计工作的过程中，我们运用职业判断，并保持职业怀疑。同时，我们也执行以下工作：", style="zero")

    paragraph13 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph13.add_run(
        "(一) 识别和评估由于舞弊或错误导致的财务报表重大错报风险，设计和实施审计程序以应对这些风险，并获取充分、适当的审计证据，作为发表审计意见的基础。由于舞弊可能涉及串通、伪造、故意遗漏、虚假陈述或凌驾于内部控制之上，未能发现由于舞弊导致的重大错报的风险高于未能发现由于错误导致的重大错报的风险。",
        style="zero")

    if context["report_params"]["internalControlAudit"]:
        paragraph14 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph14.add_run("(二) 了解与审计相关的内部控制，以设计恰当的审计程序。", style="zero")
    else:
        paragraph14 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph14.add_run("(二) 了解与审计相关的内部控制，以设计恰当的审计程序，但目的并非对内部控制的有效性发表意见。", style="zero")

    paragraph15 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph15.add_run("(三) 评价管理层选用会计政策的恰当性和作出会计估计及相关披露的合理性。", style="zero")

    four_work = ["(四) 对管理层使用持续经营假设的恰当性得出结论。同时，根据获取的审计证据，就可能导致对", companyAbbrName,
                 "持续经营能力产生重大疑虑的事项或情况是否存在重大不确定性得出结论。如果我们得出结论认为存在重大不确定性，审计准则要求我们在审计报告中提请报表使用者注意财务报表中的相关披露；如果披露不充分，我们应当发表非无保留意见。我们的结论基于截至审计报告日可获得的信息。然而，未来的事项或情况可能导致",
                 companyAbbrName, "不能持续经营。"]
    paragraph16 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph16.add_run("".join(four_work), style="zero")

    paragraph17 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph17.add_run("(五) 评价财务报表的总体列报、结构和内容，并评价财务报表是否公允反映相关交易和事项。", style="zero")

    if context["report_params"]["type"] == "合并":
        paragraph17 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph17.add_run("(六) 就{}中实体或业务活动的财务信息获取充分、适当的审计证据，以对财务报表发表审计意见。我们负责指导、监督和执行集团审计，并对审计意见承担全部责任。".format(
            context["report_params"]["CompanyAbbrName"]), style="zero")

    paragraph18 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph18.add_run("我们与治理层就计划的审计范围、时间安排和重大审计发现等事项进行沟通，包括沟通我们在审计中识别出的值得关注的内部控制缺陷。", style="zero")

    if len(context["report_params"]["keyAuditMatters"]) > 0:
        paragraph18 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph18.add_run("我们还就已遵守与独立性相关的职业道德要求向治理层提供声明，并与治理层沟通可能被合理认为影响我们独立性的所有关系和其他事项，以及相关的防范措施（如适用）。",
                            style="zero")
        paragraph18 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph18.add_run(
            "从与治理层沟通过的事项中，我们确定哪些事项对本期财务报表审计最为重要，因而构成关键审计事项。我们在审计报告中描述这些事项，除非法律法规禁止公开披露这些事项，或在极少数情形下，如果合理预期在审计报告中沟通某事项造成的负面后果超过在公众利益方面产生的益处，我们确定不应在审计报告中沟通该事项。",
            style="zero")


# 审计报告
def auditReport(context):
    companyName = context["report_params"]["companyName"]
    companyAbbrName = context["report_params"]["CompanyAbbrName"]
    reportNo = context["report_params"]["reportNo"]
    reportType = context["report_params"]["type"]
    reportDate = context["report_params"]["reportDate"]
    reportPeriod = context["report_params"]["reportPeriod"]
    accountFirms = getAccountFirmNames(context["report_params"]["accountFirm"])
    appellation = getAppellation(companyName)

    document = Document()
    # 设置中文标题

    setStyle(document)
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.add_run('审计报告', style="title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_no = document.add_paragraph()
    report_no.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    report_no.add_run(reportNo, style="zero")
    document.add_paragraph()

    appellationPara = document.add_paragraph()
    appellationPara.add_run(appellation, style="first")

    document.add_paragraph()
    paragraph1 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph1.add_run("一、审计意见", style="first")
    if reportType == "单体":
        audit_fan = ["我们审计了", companyName, "（以下简称“", companyAbbrName, "”）财务报表，包括", reportDate, "的资产负债表，", reportPeriod,
                     "的利润表、现金流量表、所有者权益变动表，以及相关财务报表附注。"]
    else:
        audit_fan = ["我们审计了", companyName, "（以下简称“", companyAbbrName, "”）财务报表，包括", reportDate, "的合并及母公司资产负债表，",
                     reportPeriod,
                     "的合并及母公司利润表、合并及母公司现金流量表、合并及母公司所有者权益变动表，以及相关财务报表附注。"]
    paragraph2 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph2.add_run("".join(audit_fan), style="zero")
    if reportType == "单体":
        audit_opinion = ["我们认为，后附的财务报表在所有重大方面按照企业会计准则的规定编制，公允反映了", companyAbbrName, reportDate, "的财务状况以及", reportPeriod,
                         "的经营成果和现金流量。"]
    else:
        audit_opinion = ["我们认为，后附的财务报表在所有重大方面按照企业会计准则的规定编制，公允反映了", companyAbbrName, reportDate, "的合并及母公司财务状况以及",
                         reportPeriod, "的合并及母公司经营成果和现金流量。"]

    paragraph3 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph3.add_run("".join(audit_opinion), style="zero")

    document.add_paragraph()
    paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph4.add_run("二、形成审计意见的基础", style="first")

    audit_basic = ["我们按照中国注册会计师审计准则的规定执行了审计工作。审计报告的“注册会计师对财务报表审计的责任”部分进一步阐述了我们在这些准则下的责任。按照中国注册会计师职业道德守则，我们独立于",
                   companyAbbrName, "，并履行了职业道德方面的其他责任。我们相信，我们获取的审计证据是充分、适当的，为发表审计意见提供了基础。"]
    paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
    paragraph5.add_run("".join(audit_basic), style="zero")

    document.add_paragraph()
    # 是否添加关键审计事项
    if len(context["report_params"]["keyAuditMatters"]) == 0:
        # 是否添加其他信息
        if context["report_params"]["otherInfo"]:
            paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph4.add_run("三、其他信息", style="first")
            addOtherInfo(document,context)

            document.add_paragraph()
            paragraph6 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph6.add_run("四、管理层和治理层对财务报表的责任", style="first")
            addManagementResponsebility(document,context)

            document.add_paragraph()
            paragraph10 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph10.add_run("五、注册会计师对财务报表审计的责任", style="first")
            addCpaResponsibility(document,context)
        else:
            paragraph6 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph6.add_run("三、管理层和治理层对财务报表的责任", style="first")
            addManagementResponsebility(document,context)

            document.add_paragraph()
            paragraph10 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph10.add_run("四、注册会计师对财务报表审计的责任", style="first")
            addCpaResponsibility(document,context)
    else:
        # 披露关键审计事项
        paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph4.add_run("三、关键审计事项", style="first")
        paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
        paragraph5.add_run("关键审计事项是我们根据职业判断，认为对本期财务报表审计最为重要的事项。这些事项的应对以对财务报表整体进行审计并形成审计意见为背景，我们不对这些事项单独发表意见。",
                           style="zero")
        for key, item in enumerate(context["report_params"]["keyAuditMatters"]):
            paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph5.add_run(item["title"], style="zero")
            paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph5.add_run("1. 事项描述", style="zero")
            for desc in item["desc"]:
                paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
                paragraph5.add_run(desc, style="zero")
            paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph5.add_run("2. 审计应对", style="zero")
            for response in item["auditResponse"]:
                paragraph5 = document.add_paragraph(style="paragraphAfterSpace")
                paragraph5.add_run(response, style="zero")
        # 是否添加其他信息
        if context["report_params"]["otherInfo"]:
            document.add_paragraph()
            paragraph4 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph4.add_run("四、其他信息", style="first")
            addOtherInfo(document,context)

            document.add_paragraph()
            paragraph6 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph6.add_run("五、管理层和治理层对财务报表的责任", style="first")
            addManagementResponsebility(document,context)

            document.add_paragraph()
            paragraph10 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph10.add_run("六、注册会计师对财务报表审计的责任", style="first")
            addCpaResponsibility(document,context)
        else:
            document.add_paragraph()
            paragraph6 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph6.add_run("四、管理层和治理层对财务报表的责任", style="first")
            addManagementResponsebility(document,context)

            document.add_paragraph()
            paragraph10 = document.add_paragraph(style="paragraphAfterSpace")
            paragraph10.add_run("五、注册会计师对财务报表审计的责任", style="first")
            addCpaResponsibility(document,context)

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()

    table = document.add_table(rows=3, cols=2)
    cell00 = table.cell(0, 0).add_paragraph()
    cell00.add_run(accountFirms[0], style="zero")
    cell00.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if len(accountFirms) > 1:
        cell00 = table.cell(0, 0).add_paragraph()
        cell00.add_run(accountFirms[1], style="zero")
        cell00.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    cell01 = table.cell(0, 1).add_paragraph()
    cell01.add_run("中国注册会计师：", style="zero")

    cell00 = table.cell(1, 0).add_paragraph()
    cell00.add_run(context["report_params"]["accountFirmAddr"], style="zero")
    cell00.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    cell11 = table.cell(1, 1).add_paragraph()
    cell11.add_run("中国注册会计师：", style="zero")

    cell21 = table.cell(2, 1).add_paragraph()
    cell21.add_run(context["report_params"]["issuanceDate"], style="zero")

    document.add_section()
    # 给报告添加页码
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document


if __name__ == '__main__':
    from project.data import testcontext
    auditReport(testcontext)
