import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START


# 判断字符串是否为数字
def is_number(s):
    try:
        float(s)
        return True
    except Exception:
        pass
    return False


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run(style="zero")
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = '第 '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run(style="zero")
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' 页 共 '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

    end_run = paragraph.add_run(style="zero")
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' 页'
    end_run._r.append(t2)


# 设置单元格边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


_MAPPING = ('零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九')
_P0 = ('', '十', '百', '千',)
_S4 = 10 ** 4


def to_chinese(num):
    if isinstance(num, str):
        return 0
    if num < 20:
        return _MAPPING[num]
    else:
        lst = []
        while num >= 10:
            lst.append(num % 10)
            num = num / 10
        lst.append(num)
        # c = len(lst)  # 位数
        result = ''

        for idx, val in enumerate(lst):
            val = int(val)
            if val != 0:
                result += _P0[idx] + _MAPPING[val]
        return result[::-1]


# 添加标题头
def addTitle(document, title, level, indent):
    basic_title = document.add_heading("", level=level)
    if level >= 3:
        basic_title.add_run(title, style="first")
    else:
        basic_title.add_run(title, style="first").bold = True
    basic_title.style.paragraph_format.line_spacing = 1.5
    basic_title.style.paragraph_format.space_before = Pt(0)
    if indent:
        basic_title.style.paragraph_format.first_line_indent = Pt(24)


# 添加段落
def addParagraph(document, content, paragraphStyle):
    paragraph = document.add_paragraph(style=paragraphStyle)
    paragraph.add_run(content, style="zero")


# 添加表格边框
def createBorderedTable(document, rowLength, columnLength, innerLine="dotted"):
    table = document.add_table(rowLength, columnLength)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"sz": 12, "val": "single", "space": "0"})
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 12, "val": "single", "space": "0"})
    for row in table.rows[0:len(table.rows) - 1]:
        for cell in row.cells:
            set_cell_border(cell, bottom={"sz": 6, "val": innerLine, "space": "0"})
    for key, column in enumerate(table.columns):
        if key == len(table.columns) - 1:
            continue
        for cell in column.cells:
            set_cell_border(cell, right={"sz": 6, "val": innerLine, "space": "0"})
    # 设置表格格式
    table.autofit = True
    return table


# 设置单元格格式

def setCell(cell, cellText, alignment, toFloat=True, style="tableCharacter"):
    if is_number(cellText) and toFloat:
        f = float(cellText)
        if abs(f)<1e-7:
            cellText=""
        else:
            cellText = '{:,.2f}'.format(float(cellText))
    cellString = str(cellText)
    if cellString == "nan":
        cellString = ""

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    paragraph_format = p.paragraph_format
    # paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(1)
    paragraph_format.space_before = Pt(1)
    paragraph_format.alignment = alignment
    p.add_run(str(cellString), style=style)


def setNotToFloatCell(title, cell, j, row, alignment):
    if title[j] in ["级次", "序号"]:
        setCell(cell, row[j], alignment, toFloat=False)
    else:
        setCell(cell, row[j], alignment, toFloat=True)


# 向表格中添加数据
'''
style:
1:单元格全部左对齐
2：标题居中，第一列除最后单元格居中外其余全部左对齐，其他列右对齐
3：标题居中，第一列左对齐，其余右对齐
4:全部居中
5：标题居中，其余左对齐
6：数字靠右，文字靠左，标题居中，第一列靠左
'''


# 仅适用于单行标题
def addTable(document, table, style=1):
    # 获取标题和单元格
    title = table["columns"]
    cells = table["data"]
    if len(cells) == 0:
        addParagraph(document, "不适用", "paragraph")
        return
    columnLength = len(cells[0])
    # 没有标题
    if len(title) == 0:
        data = cells
        rowLength = len(cells)
        table = createBorderedTable(document, rowLength, columnLength)
        if style == 1:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif style == 2:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if j == 0 and i != rowLength - 1:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                    elif j == 0 and i == rowLength - 1:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.CENTER)
                    else:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
        elif style == 3:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if j == 0:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                    else:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
        elif style == 4:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.CENTER)
        elif style == 5:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif style == 6:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if j == 0:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                    else:
                        if is_number(row[j]):
                            setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
                        else:
                            setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
    else:
        data = [title, *cells]
        rowLength = len(cells) + 1
        table = createBorderedTable(document, rowLength, columnLength)
        if style == 1:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif style == 2:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if i == 0:
                        setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    else:
                        if j == 0 and i != rowLength - 1:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)
                        elif j == 0 and i == rowLength - 1:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
                        else:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.RIGHT)
        elif style == 3:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if i == 0:
                        setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    else:
                        if j == 0:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)
                        else:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.RIGHT)
        elif style == 4:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
        elif style == 5:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if i == 0:
                        setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    else:
                        setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif style == 6:
            for i, row in enumerate(data):
                for j, cell in enumerate(table.rows[i].cells):
                    if i == 0:
                        setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    else:
                        if j == 0:
                            setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)
                        else:
                            if is_number(row[j]):
                                setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.RIGHT)
                            else:
                                setNotToFloatCell(title, cell, j, row, WD_PARAGRAPH_ALIGNMENT.LEFT)


# 添加众向合并表格，自动合并下一个为空的单元格
def addCombineTableContent(table, dc, titleLength):
    for key in range(len(dc["index"])):
        row = key + titleLength
        for j in range(len(dc["columns"])):
            cell = table.cell(row, j)
            if str(dc["data"][key][j]) == "nan":
                cell.merge(table.cell(row - 1, j))
            else:
                setCell(cell, str(dc["data"][key][j]), WD_PARAGRAPH_ALIGNMENT.LEFT)


# 添加众向合并表格，自动合并下一个为空的单元格
def addContentToCombineTitle(document, dc, table, titleLength, style=1):
    # 获取标题和单元格
    cells = dc["data"]
    if len(cells) == 0:
        addParagraph(document, "不适用", "paragraph")
        return
    # 没有标题
    data = cells
    rowLength = len(cells) + titleLength
    if style == 1:
        for i, row in enumerate(data):
            i = i + titleLength
            for j, cell in enumerate(table.rows[i].cells):
                setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
    elif style == 2:
        for i, row in enumerate(data):
            i = i + titleLength
            for j, cell in enumerate(table.rows[i].cells):
                if j == 0 and i != rowLength - 1:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                elif j == 0 and i == rowLength - 1:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.CENTER)
                else:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
    elif style == 3:
        for i, row in enumerate(data):
            i = i + titleLength
            for j, cell in enumerate(table.rows[i].cells):
                if j == 0:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                else:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
    elif style == 4:
        for i, row in enumerate(data):
            i = i + titleLength
            for j, cell in enumerate(table.rows[i].cells):
                setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.CENTER)
    elif style == 5:
        for i, row in enumerate(data):
            for j, cell in enumerate(table.rows[i].cells):
                setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
    elif style == 6:
        for i, row in enumerate(data):
            i = i + titleLength
            for j, cell in enumerate(table.rows[i].cells):
                if j == 0:
                    setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)
                else:
                    if is_number(row[j]):
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.RIGHT)
                    else:
                        setCell(cell, row[j], WD_PARAGRAPH_ALIGNMENT.LEFT)


# 检查标题左面是否全部为空
def checkLeftSpace(row, j):
    for i in range(j):
        if row[i] != "nan":
            return False
    return True

# # 检查标题左面是否全部为空
def checkTableLeftSpace(row, j):
    for i in range(j):
        if row[i] != "":
            return False
    return True


# 添加合并标题
# 左侧单元格空格众向合并，
def addCombineTableTitle(table, titles):
    for i, row in enumerate(titles):
        for j, cellText in enumerate(row):
            cell = table.cell(i, j)
            if cellText == "nan":
                if j >= 1:
                    if checkLeftSpace(row, j):
                        cell.merge(table.cell(i - 1, j))
                    else:
                        cell.merge(table.cell(i, j - 1))
                else:
                    cell.merge(table.cell(i - 1, j))
            else:
                setCell(cell, cellText, WD_PARAGRAPH_ALIGNMENT.CENTER)
    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})


# 添加横向内容
'''
type:合并、单体
'''


def addLandscapeContent(document, func, *args):
    # 设置横向
    section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    section.orientation = WD_ORIENTATION.LANDSCAPE
    page_h, page_w = section.page_width, section.page_height
    section.page_width = page_w
    section.page_height = page_h

    # 添加内容
    func(document, *args)

    # 重新设置为众向
    section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    section.orientation = WD_ORIENTATION.PORTRAIT
    page_h, page_w = section.page_width, section.page_height
    section.page_width = page_w
    section.page_height = page_h


# 替换名称中的特殊符号☆,△和空格
def handleName(name):
    name = name.replace("☆", "")
    name = name.replace("△", "")
    name = name.replace("*", "")
    name = name.replace("#", "")
    name = name.replace(" ", "")
    return name


# 通过报表名称搜索对应的报表项目(资产负债表/利润表/现金流量表，hasNum表示是否仅搜索有数据的项目
def searchRecordItemByName(name, records, fillNum=False):
    for record in records:
        if fillNum:
            if handleName(record["name"]) == name and record["fillNum"]:
                return record
        else:
            if handleName(record["name"]) == name:
                return record
    return None


# 查找模型
def searchModel(companyType, fsType, table, comparativeTable):
    for item in comparativeTable:
        if item["companyType"] == companyType and item["fsType"] == fsType and item["table"] == table:
            return item["model"]
    return None


# 获取报告附注编码
def getNoteNum(context,isParent):
    companyType = context["report_params"]["companyType"]
    # 报告类型
    reportType = context["report_params"]["type"]

    if isParent:
        if companyType == "国有企业":
            if reportType == "合并":
                return "十二"
            else:
                return "十一"
        else:
            if reportType == "合并":
                if context["noteAppend"]["shareBasedPayment"]:
                    return "十六"
                else:
                    return "十五"
            else:
                if context["noteAppend"]["shareBasedPayment"]:
                    return "十四"
                else:
                    return "十三"
    else:
        if companyType == "国有企业":
            if reportType == "合并":
                return "八"
            else:
                return "七"
        else:
            if reportType == "合并":
                return "六"
            else:
                return "六"

# 是否向上合并单元格
def isCombine(i, j, combineCells):
    for cell in combineCells:
        if i == cell[0] and j == cell[1]:
            return True
    return False

# 为表格添加合并标题，指定列向上合并
def addCombineTitleSpecialReceivable(titles, table, combineCells):
    for i, row in enumerate(titles):
        for j, cellText in enumerate(row):
            cell = table.cell(i, j)
            if cellText == "nan":
                # 最后一列，最后一行
                if isCombine(i, j, combineCells):
                    cell.merge(table.cell(i - 1, j))
                    continue
                if j >= 1:
                    if checkLeftSpace(row, j):
                        cell.merge(table.cell(i - 1, j))
                    else:
                        cell.merge(table.cell(i, j - 1))
                else:
                    cell.merge(table.cell(i - 1, j))
            else:
                setCell(cell, cellText, WD_PARAGRAPH_ALIGNMENT.CENTER)
    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})

# 替换表格内容空白回车
def replace_doc_table(titles,table):
    for key,row in enumerate(table.rows):  # 遍历表格中的所有行
        if key> len(titles)-1:
            return
        for cell in row.cells:  # 遍历行中的所有单元格
            newText = cell.text.replace('\n','')
            cell.text = ""
            setCell(cell, newText, WD_PARAGRAPH_ALIGNMENT.CENTER, False, "tableSmallerCharacter")


# 为表格添加合并标题，指定列向上合并,合并所有者权益变动表表头
def combineTitles(titles, table, combineCells,lastLine):
    for i, row in enumerate(titles):
        for j, cellText in enumerate(row):
            cell = table.cell(i, j)
            if cellText == "":
                if lastLine and i==len(titles)-1:
                    try:
                        cell.merge(table.cell(i - 1, j))
                    except Exception as e:
                        print("001合并单元格错误", e, i, j)
                # 最后一列，最后一行
                if isCombine(i, j, combineCells):
                    try:
                        cell.merge(table.cell(i - 1, j))
                    except Exception as e:
                        print("002合并单元格错误", e, i, j)
                    continue
                if j >1:
                    if checkTableLeftSpace(row, j):
                        try:
                            cell.merge(table.cell(i - 1, j))
                        except Exception as e:
                            print("003合并单元格错误",e,i,j)
                    else:
                        try:
                            cell.merge(table.cell(i, j - 1))
                        except Exception as e:
                            print("004合并单元格错误", e, i, j)
                else:
                    try:
                        cell.merge(table.cell(i - 1, j))
                    except Exception as e:
                        print("005合并单元格错误", e, i, j)
    replace_doc_table(titles,table)
    set_cell_border(table.cell(0, len(table.columns) - 1), right={"sz": 0, "val": "", "space": "0"})

# 排序
# 合计最后一行，其他项目倒数第二行，其他项目按照column_name的从大到小排列
def df_sort(df,index_name="项目",column_name="本期数",last_names=("其他","合计")):
    '''

    :param df: 要排序的df
    :param index_name: 条件列，如项目
    :param column_name: 值列，从大到小排列
    :param last_names: 条件列中不参与排序的值
    :return: 排序后的结果
    '''
    last_condition = False
    first_conditon = True
    try:
        for name in last_names:
            last_condition = last_condition | df[index_name].str.contains(name)
            first_conditon = first_conditon & (~(df[index_name].str.contains(name)))
    except Exception as e:
        return df
    df_last = df[last_condition]
    df_first = df[first_conditon]
    df_first = df_first.sort_values(by=column_name, ascending=False)
    df_all = pd.concat([df_first, df_last], ignore_index=True)
    return df_all

# 根据条件过滤df
def filterDateFrame(sheetName, xlsxPath,conditions=("期末数","期初数")):
    df = pd.read_excel(xlsxPath, sheet_name=sheetName)
    if len(conditions)==0:
        return df
    s = False
    values_dict = dict()
    for condition in conditions:
        values_dict[condition] = 0.00
    df = df.fillna(value=values_dict)
    for condition in conditions:
        try:
            s = s | (df[condition].abs() > 0)
        except Exception as e:
            print(e)
            return df
    df1 = df[s]
    return df1
