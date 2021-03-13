import os
from docx import Document
from win32com.client import Dispatch

# 添加文本框
def addTextBox(docFile,content):
    app = Dispatch("Word.Application")
    doc = app.Documents.Open(docFile)

    tb = doc.Shapes.AddTextbox(1, 339, 28, 162, 78)
    content_pg  = tb.TextFrame.TextRange.Paragraphs.Add()
    content_pg.Range.Font.Name, content_pg.Range.Font.Size = '黑体', 7
    content_pg.Range.ParagraphFormat.Alignment = 0
    content_pg.Range.InsertBefore(content)

    # tb.TextFrame.TextRange.Text = "Hello^pHello"
    # content_pg = tb.TextFrame.TextRange.Paragraphs.Add()
    # content_pg.Text = "Hello^pHello"
    tb.TextFrame.MarginTop = 0
    tb.TextFrame.MarginLeft = 0
    tb.Fill.Visible = 0
    tb.Line.Visible = 0

    doc.SaveAs2(docFile)
    doc.Close()
    app.Application.Quit()

# 替换文档前几段中的文字
def replaceDoc(document,old_text,new_text,paragraphsNum):
    for paragraph in document.paragraphs[:paragraphsNum]:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text=run.text.replace(old_text,new_text)

# 添加页眉图片
def addHeader(document,content):
    # 获取第一个节的页眉
    header = document.sections[0].header
    # 获取页眉的第一个段落
    paragraph = header.paragraphs[0]
    # 添加图片
    run = paragraph.add_run()
    run.add_picture(content)


def main(path):
    allow_suffix = ".docx"
    content = "亚太（集团）会计师事务所（特殊普通合伙）\n中国北京丰台区丽泽路16号院3号楼20层2001\r\n邮编100004\r\n电话 +86 10 88312386\r\n传真 +86 10 88386116\r\nwww.apag-cn.com"
    oldText = "董事会"
    newText = "全体股东"
    headerImg = "1.png"

    all_files = [f for f in os.listdir(path) if f.endswith(allow_suffix)]
    for file in all_files:
        if file.startswith("~$"):
            continue
        filepath = os.path.join(path, file)
        document = Document(filepath)
        # 添加表头
        addHeader(document, headerImg)
        replaceDoc(document,oldText , newText, 5)
        document.save(filepath)
        addTextBox(filepath,content)



if __name__ == '__main__':
    path = r"D:\审计\我的文件2021\发展集团2020年审\发展集团2020TB\审计报告"
    main(path)